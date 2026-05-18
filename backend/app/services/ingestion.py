import hashlib
import io
import uuid
from pathlib import Path
from typing import Optional

import tiktoken
from openai import AsyncOpenAI
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.logging import get_logger
from app.models.database import Document, DocumentChunk

_settings = get_settings()
_logger = get_logger(__name__)
_enc = tiktoken.get_encoding("cl100k_base")


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _extract_text_pdf(data: bytes) -> list[tuple[str, int]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[tuple[str, int]] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((text, i))
    return pages


def _extract_text_docx(data: bytes) -> list[tuple[str, int]]:
    import docx

    doc = docx.Document(io.BytesIO(data))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [(full_text, 0)]


def _extract_text_txt(data: bytes) -> list[tuple[str, int]]:
    return [(data.decode("utf-8", errors="replace"), 0)]


def _extract_text(data: bytes, file_type: str) -> list[tuple[str, int]]:
    extractors = {
        "pdf": _extract_text_pdf,
        "docx": _extract_text_docx,
        "txt": _extract_text_txt,
    }
    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    return extractor(data)


def _count_tokens(text: str) -> int:
    return len(_enc.encode(text))


def _chunk_text(
    text: str,
    chunk_size: int = 512,
    overlap: int = 64,
) -> list[str]:
    """Recursive character splitting by tokens."""
    separators = ["\n\n", "\n", ". ", " ", ""]
    tokens = _enc.encode(text)

    if len(tokens) <= chunk_size:
        return [text] if text.strip() else []

    chunks: list[str] = []
    start = 0
    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = _enc.decode(chunk_tokens)

        # Try to find a natural break point near the end
        if end < len(tokens):
            for sep in separators[:-1]:
                idx = chunk_text.rfind(sep, len(chunk_text) // 2)
                if idx != -1:
                    chunk_text = chunk_text[: idx + len(sep)]
                    chunk_tokens = _enc.encode(chunk_text)
                    end = start + len(chunk_tokens)
                    break

        if chunk_text.strip():
            chunks.append(chunk_text.strip())
        start = end - overlap

    return chunks


async def _embed_texts(texts: list[str], client: AsyncOpenAI) -> list[list[float]]:
    response = await client.embeddings.create(
        model=_settings.embedding_model,
        input=texts,
    )
    return [item.embedding for item in response.data]


class IngestionService:
    def __init__(self) -> None:
        self._openai = AsyncOpenAI(api_key=_settings.openai_api_key)

    async def ingest(
        self,
        data: bytes,
        filename: str,
        department: Optional[str],
        db: AsyncSession,
    ) -> Document:
        file_type = Path(filename).suffix.lstrip(".").lower()
        file_hash = _hash_bytes(data)

        existing = await db.scalar(select(Document).where(Document.file_hash == file_hash))
        if existing:
            _logger.info("document_already_exists", extra={"filename": filename, "hash": file_hash})
            return existing

        pages = _extract_text(data, file_type)
        all_chunks: list[tuple[str, int]] = []
        for page_text, page_num in pages:
            for chunk in _chunk_text(page_text, _settings.chunk_size, _settings.chunk_overlap):
                all_chunks.append((chunk, page_num))

        _logger.info("chunks_created", extra={"filename": filename, "count": len(all_chunks)})

        doc = Document(
            id=str(uuid.uuid4()),
            filename=filename,
            file_hash=file_hash,
            department=department,
            file_type=file_type,
            total_chunks=len(all_chunks),
        )
        db.add(doc)
        await db.flush()

        # Embed in batches of 100
        batch_size = 100
        for batch_start in range(0, len(all_chunks), batch_size):
            batch = all_chunks[batch_start : batch_start + batch_size]
            texts = [c[0] for c in batch]
            embeddings = await _embed_texts(texts, self._openai)

            for i, (chunk_text, page_num) in enumerate(batch):
                db.add(
                    DocumentChunk(
                        id=str(uuid.uuid4()),
                        document_id=doc.id,
                        content=chunk_text,
                        embedding=embeddings[i],
                        chunk_index=batch_start + i,
                        page_number=page_num if page_num else None,
                        token_count=_count_tokens(chunk_text),
                    )
                )

        await db.commit()
        await db.refresh(doc)
        _logger.info("document_ingested", extra={"document_id": doc.id, "chunks": len(all_chunks)})
        return doc
